import docx

class DocxEditor:
    @staticmethod
    def extract_bullets(docx_path):
        doc = docx.Document(docx_path)
        bullets = []
        for p in doc.paragraphs:
            text = p.text.strip()
            if len(text) > 15:
                bullets.append(text)
        return bullets
        
    @staticmethod
    def _is_match(text1, text2, threshold=0.9):
        """Helper to check if two strings match after basic normalization or similarity."""
        if not text1 or not text2: return False
        n1 = " ".join(text1.lower().split()).strip().strip('.')
        n2 = " ".join(text2.lower().split()).strip().strip('.')
        
        if n1 == n2: return True
        
        # Fallback to SequenceMatcher for minor AI-introduced wording drifts
        import difflib
        ratio = difflib.SequenceMatcher(None, n1, n2).ratio()
        return ratio >= threshold

    @staticmethod
    def apply_edits(in_path, out_path, edits):
        doc = docx.Document(in_path)
        from docx.enum.text import WD_COLOR_INDEX
        
        rep_count = 0
        deleted_count = 0
        
        # Convert edits to list of tuples for ordered processing or multi-matching
        edit_list = [(e['old'], e['new']) for e in edits]
        
        # Track which edits have been satisfied to avoid double-processing
        satisfied_indices = set()
        
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text: continue
            
            for idx, (old_text, new_text) in enumerate(edit_list):
                if idx in satisfied_indices: continue
                
                if DocxEditor._is_match(text, old_text):
                    satisfied_indices.add(idx)
                    
                    if new_text == "":
                        # Delete the paragraph effectively
                        p._element.getparent().remove(p._element)
                        deleted_count += 1
                    else:
                        # Capture style from original runs
                        font_name = None
                        font_size = None
                        is_bold = None
                        
                        if len(p.runs) > 0:
                            first_font = p.runs[0].font
                            font_name = first_font.name
                            font_size = first_font.size
                            is_bold = first_font.bold
                        
                        # Apply new text
                        p.text = new_text
                        
                        # Re-apply style and highlight
                        if len(p.runs) > 0:
                            if font_name: p.runs[0].font.name = font_name
                            if font_size: p.runs[0].font.size = font_size
                            if is_bold is not None: p.runs[0].font.bold = is_bold
                            
                            # Apply YELLOW HIGHLIGHT to tailored content
                            p.runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW
                    
                    rep_count += 1
                    break # Move to next paragraph
                
        print(f"DOCX Parser natively reformatted {rep_count - deleted_count} semantic blocks and securely DELETED {deleted_count} weak bullets for 1-Page constraints!")
        doc.save(out_path)
        return out_path
        
    @staticmethod
    def create_cover_letter_docx(text, output_path):
        doc = docx.Document()
        
        # Enforce exact Roboto typography standard natively identical to Resume baseline cleanly
        style = doc.styles['Normal']
        style.font.name = 'Roboto'
        
        # Enforce heading cascades
        for i in range(1, 10):
            if f'Heading {i}' in doc.styles:
                doc.styles[f'Heading {i}'].font.name = 'Roboto'
        
        lines = text.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            if line.startswith('#'):
                level = min(line.count('#'), 9)
                clean_text = line.replace('#', '').strip()
                doc.add_heading(clean_text, level=level if level > 0 else 1)
            else:
                doc.add_paragraph(line)
                
        doc.save(output_path)
        return output_path
